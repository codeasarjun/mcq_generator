from io import BytesIO
from docx import Document

def create_word(questions):
    output = BytesIO()
    doc = Document()
    doc.add_heading('Generated Questions', 0)

    for q in questions:
        doc.add_paragraph(f"Question: {q['ques']}")
        
        for letter_key, choice in q['choices'].items():
            doc.add_paragraph(f"{letter_key}. {choice.capitalize()}")

        doc.add_paragraph(f"Correct Answer: {q['correct_ans'].capitalize()}")
        doc.add_paragraph()  # Blank line between questions

    doc.save(output)
    output.seek(0)
    return output
